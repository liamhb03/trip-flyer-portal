#!/usr/bin/env python3
"""Apply todo.md revisions to the SAF literature review docx."""

import copy
import io
import re
import shutil
import zipfile
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

SRC = "/Users/liamhornschild-bear/docchanges/SAF_Lit_Review_Draft_2_REVISED (1)_orig.docx"
OUT = "/Users/liamhornschild-bear/docchanges/SAF_Lit_Review_Draft_2_REVISED (1) (1).docx"

OOXML_PARTS_TO_RESTORE = (
    "word/styles.xml",
    "word/numbering.xml",
    "word/settings.xml",
)

INTRO = """Unlike most industrial sectors, commercial aviation has never had to rebuild its energy base. Since the birth of commercial flight, the same class of fuel, kerosene refined from crude oil, has powered essentially every aircraft in the sky. This is not simply a matter of habit or an inherited supply chain. Aircraft engines, airport fuelling infrastructure, and international certification regimes are all built around a narrow specification window, and any substitute fuel must slot into that window rather than force the industry to redesign around it. That constraint has kept aviation's energy base essentially unchanged even as other heavy industries, from utilities to shipping, have diversified toward multiple energy carriers over the past two decades.

This structural rigidity is now colliding with a wave of binding European climate policy. Mandates such as the ReFuelEU Aviation Regulation and the phase-out of free allowances under the EU Emissions Trading System are forcing a shift toward sustainable aviation fuel (SAF) on a fixed timetable, regardless of whether the underlying supply chain, cost structure, or production capacity is ready to support it. The full detail of this regulatory architecture, and how it interacts with Dutch national policy, is set out in Section 3, immediately below. What matters for the argument here is more basic: the mandate does not, by itself, create a viable supply chain, and a supply chain is not the same thing as a business model. Bidmon and Knab (2018) describe the pre-existing arrangement among fuel producers, distributors, airports, and airlines as an industry recipe, a collectively stabilised logic of value creation and capture that has held for decades. Regulatory pressure is now placing that recipe under direct challenge, but it says nothing about what should replace it, who should build it, or how the costs and risks of the transition should be shared.

This paper examines that gap through a specific circular innovation pathway: waste-to-SAF via hydrothermal liquefaction (HTL), a technology that converts organic waste streams into a biocrude intermediate that must then be upgraded into jet-range molecules. A pilot collaboration between the University of Amsterdam and Schiphol Airport has demonstrated that HTL can produce fuel-grade output from airport waste at lab scale. What it has not demonstrated, and cannot demonstrate through chemistry alone, is that this pathway can become a commercially viable operation. Technical feasibility and business viability are different questions, and the gap between them is not primarily a technology problem. It is an organisational and inter-organisational one: who invests, who bears the risk if a batch fails to meet specification, who certifies the output, and who is willing to pay a premium for it.

The research question that follows from this is: how does a policy-driven innovation get translated into a viable, sustainable business model within the aviation ecosystem? This breaks down into five more specific questions. First, what does the current aviation fuel business model look like, and where do value and risk currently sit within it? Second, what structural elements of that model have to change to accommodate a waste-to-SAF pathway? Third, which actors have to participate in that change, including actors that play no role in the conventional fuel chain today? Fourth, how are the resulting costs and risks likely to be redistributed across the value chain? And fifth, can this transition happen without breaking the operational continuity the sector depends on, or does it necessarily introduce friction and disruption along the way?

Answering these questions draws on three sources that have so far developed largely apart from one another: a review of the business model innovation and sustainable and circular business model literature, a mapping of the regulatory instruments governing SAF and waste valorisation at EU and Dutch level, and early interview material gathered from across the Dutch aviation fuel ecosystem. Section 2 sets out the methodology behind all three. Section 3 covers the regulatory architecture in detail. Section 4 reviews the business model innovation literature. Sections 5 through 8 build toward an integrated analytical framework drawing on the ecosystem, coalition formation, and financial viability literatures. Section 9 turns to the interview material. Section 10 discusses what the literature and the interviews together tell us. Section 11 concludes."""

POLICY_OPENING = """Before turning to the business model literature, this section lays out the regulatory architecture that the rest of the paper treats as a fixed constraint. It covers the binding EU-level instruments (ReFuelEU Aviation, RED III, the EU ETS, the Waste Framework Directive, and the EU Taxonomy Regulation), the certification and audit rules that convert physical feedstock into a compliance claim, and the Dutch national instruments that sit on top of the EU framework, including the Klimaatwet, the Luchtvaartnota, and the Circulair Materialenplan. The detail here is deliberately dense, since it is the fixed input the rest of the analysis works from, not itself the object of analysis. Section 4 picks up the theoretical question of how business models respond to and reshape instruments of this kind."""

METHODOLOGY = """This review combines three strands of work: a systematic literature search on business model innovation and sustainable business model theory, a targeted regulatory mapping of the instruments governing waste-to-SAF pathways in the EU and the Netherlands, and a set of semi-structured interviews with stakeholders positioned at different points of the Dutch aviation fuel value chain. The literature and policy work ground the analytical framework developed in Sections 4 through 8. The interviews, reported in Section 9, provide early empirical evidence against which that framework is read in the discussion in Section 10."""

METH_21 = """Two independent databases were searched, Business Source Premier (EBSCO) and Scopus, covering the period from database inception through May 2026. Search terms were structured using a Venn diagram logic with two circles, an initial design that separated a theoretical circle covering business model and sustainable business model literature from a phenomenon circle covering aviation sustainability literature. The phenomenon circle was later removed from the combined query in consultation with a research team member, since it introduced excessive noise from aeronautical engineering literature, and aviation-specific terms were instead folded into a dedicated intersection query. The theoretical circle (Circle 1) combined a block of business model innovation and transformation terms with a block of sustainability and circular economy terms, searched across subject headings, titles, abstracts, and author keywords in EBSCO, and across title, abstract, and keyword fields in Scopus. The intersection query (Sustainable Business Models in Aviation) combined four blocks covering business model terminology, fuel and feedstock terminology, aviation terminology, and sustainability and decarbonisation terminology. Wildcard truncation and proximity operators were used to capture term variants, for example word-distance operators of three words to capture phrases such as emissions reduction or carbon minimisation. Filters were applied for English language, peer-reviewed status, and academic journal or article document type. Record counts were logged at each stage, with Business Source Premier Circle 1 returning 2,526 records, Business Source Premier Intersection returning 29 records, Scopus Circle 1 returning 2,274 records, and Scopus Intersection returning 341 records."""

METH_22 = """All records were exported to Zotero for reference management. Manual deduplication in Zotero proved too slow given the volume of records, so an automated deduplication tool (zoplicate) was used instead. The combined pre-deduplication set totalled approximately 5,170 records, reduced to 4,301 records after deduplication. The deduplicated set was uploaded to ASReview for screening, with title coverage at 100 percent, abstract coverage at 99 percent, and DOI or URL coverage at 92.6 percent. Records were screened and labelled as relevant, not relevant, or left unlabelled. The ABS Academic Journal Guide 2024 rankings were then applied, retaining only journals classified in categories 1 and 2, which reduced the set to 224 papers. A quality check was performed on the next 200 records in the ASReview ranking, continuing until 50 consecutive irrelevant records were reached, followed by team discussion of borderline inclusions."""

METH_23 = """Backward and forward reference checking was conducted on key included articles, following citation chains to and from foundational sources such as Foss and Saebi (2017) and Pieroni et al. (2019). This step surfaced additional articles not captured by the original search strings, which were then screened using the same relevance criteria."""

METH_24 = """Because the case study sits at the intersection of business model theory and a dense regulatory environment, a separate mapping exercise was conducted for the legal and policy instruments governing SAF and waste valorisation. This mapping followed a trickle-down approach, starting from the ReFuelEU Aviation Regulation as the anchor instrument. Each operative article of ReFuelEU was read in full on EUR-Lex, and cross-references within the text were followed outward to related instruments. This process led from ReFuelEU to RED III for feedstock eligibility, to the EU ETS Directive and its Monitoring and Reporting Regulation for carbon accounting, to the Waste Framework Directive for waste classification, and to the EU Taxonomy Regulation for sustainable finance criteria. Implementing and delegated acts referenced within these primary instruments were retrieved and reviewed in the same way. The same trickle-down logic was applied at the Dutch level, starting from national instruments that transpose or sit alongside the EU rules, drawing on wetten.overheid.nl for statutory texts and open.overheid.nl and relevant ministry websites for policy documents. Where a Dutch instrument referenced an EU directive or regulation, or vice versa, that link was followed to confirm how the national instrument transposes or operationalises the EU-level rule. This citation-following approach was chosen over a keyword search of legal databases because the regulatory architecture is deeply cross-referential, and a keyword search risked missing implementing detail held in cross-referenced articles rather than in the primary instrument itself."""

METH_25 = """Alongside the literature and regulatory mapping, five semi-structured interviews were conducted with stakeholders positioned at different points of the aviation fuel value chain, following the interview guide developed for the ENLENS project (reproduced in the appendix). Participants were selected to cover distinct functional positions in the value chain rather than for statistical representativeness: an infrastructure and logistics actor, an airline fuel buyer, a former employee of a SAF producer working specifically on the voluntary certificate market, a technology supplier to SAF producers, and a major integrated energy company active in both SAF production and trading. The participants are summarised in the table below. Interviews were conducted between April and June 2026, lasted between 20 and 55 minutes, and were recorded and transcribed with participant consent under the informed consent procedure described in the project's information sheet. All participants were offered pseudonymisation; where this document names a participant's employer, that reflects the participant's own choice on the consent form or a role so specific to one organisation (for example, board membership of a named joint venture) that anonymisation was not meaningful. Interviews are treated as illustrative rather than statistically representative, given the small sample size and the early stage of the broader ENLENS project, which anticipates further interviews with feedstock suppliers, certification bodies, and Dutch policymakers in later phases."""

LIT_TABLE_HEADERS = [
    "Concept / finding",
    "Key source(s)",
    "Core argument",
    "Business model mechanism",
    "Relevance to waste-to-SAF / HTL case",
]

OVERVIEW_TABLES = {
    2: {
        "caption": "Table 2 synthesizes the definitional building blocks reviewed in this section, moving from the business model as unit of analysis through sustainable and circular distinctions to the scale of change the waste-to-SAF pathway requires.",
        "title": "Table 2: Business model innovation and the sustainable business model, key definitional building blocks",
        "rows": [
            ("group", "Business model as unit of analysis"),
            ("data", ["Business model definition", "Teece (2010)", "A business model is the architecture of value creation, delivery, and capture", "Baseline analytical lens for the whole review", "Sets the baseline lens the whole review uses to analyse the waste-to-SAF pathway"]),
            ("data", ["Modular vs. architectural change", "Foss and Saebi (2017)", "BMI is designed, nontrivial change to a model's elements or their architecture", "Distinguishes incremental procurement from systemic redesign", "Adding SAF procurement is modular; redesigning airport waste into a circular feedstock chain is architectural, and the harder of the two"]),
            ("data", ["Fit with existing recipe over tech readiness", "Richter (2013)", "In German utilities, incumbents built viable models for large-scale renewables but not distributed generation", "Industry recipe fit may outweigh technological readiness", "Suggests fit with the existing aviation fuel recipe may matter as much as HTL's technical readiness"]),
            ("data", ["Sustainable innovation design criteria", "Boons and Ludeke-Freund (2013)", "Sustainable innovation requires ecological/social value, shared upstream accountability, a motivating customer interface, and fair cost/benefit distribution", "Normative design criteria, not descriptive frameworks", "No single firm governs feedstock or compliance value; alignment across actors is a precondition"]),
            ("data", ["BMI / SBM / SBMI distinction", "Shakeel et al. (2020)", "BMI, SBM, and SBMI differ in where novelty and sustainability co-occur across components", "Requires coordinated change across the entire business model", "Waste-to-SAF needs SBMI-level coordination, not a single-component tweak"]),
            ("data", ["Ecosystem-level gap in BMI research", "Zhang et al. (2024)", "Review of 1,032 BMI publications finds growing interest in sustainable/circular BMI but limited ecosystem-level analysis", "Names a systematic gap in the literature", "Directly names the gap this paper's ecosystem-spanning waste-to-SAF case is positioned to help fill"]),
            ("group", "Sustainable vs. circular business models"),
            ("data", ["CBM as subcategory of SBM (contested)", "Geissdoerfer et al. (2020)", "Circular business models focus on cycling, extending, or dematerialising loops; whether this is automatically sustainable is contested", "Circular logic is necessary but not sufficient for sustainability", "Closing the waste loop via HTL is necessary but not sufficient for genuine sustainability"]),
            ("data", ["Ecological modernist blind spot", "Hofmann (2019)", "Most CBM research decouples growth from resource use without questioning growth itself, and underweights rebound effects", "Rebound and absolute emissions remain under-examined", "If waste-to-SAF simply enables more flying, absolute emissions could still rise even as the loop closes"]),
            ("data", ["CE vs. sustainability-oriented BMI split", "Pieroni et al. (2019)", "CE-oriented work centres resource efficiency and growth; sustainability-oriented work centres social impact and long-term viability; only ~20% address implementation", "Design-implementation gap is institutional, not just managerial", "The design-implementation gap in waste-to-fuel is an institutional and financial gap, not just a modelling weakness"]),
            ("group", "Scale of required change"),
            ("data", ["Barrier / facilitator / niche innovator", "Bidmon and Knab (2018)", "Business models in sustainability transitions act as barriers, facilitators, or independent niche innovations", "Frames how pre-commercial pathways become commercial", "Frames waste-to-SAF as a niche innovation that needs a business model built around it to become commercially viable"]),
            ("data", ["Paths for incumbents under pressure", "Nguyen and Mori (2026)", "Oil and gas incumbents typically double down, diversify, or (rarely) go all-in on the alternative; transformation is enabled by scale, ownership, and asset-age conditions", "Most actors diversify minimally under contradictory policy signals", "Most airports/airlines will diversify minimally; Schiphol's HTL pathway asks some actors to absorb the cost of the more radical path"]),
        ],
    },
    3: {
        "caption": "Table 3 synthesizes the policy-business model co-evolution literature reviewed in this section, from macro-level policy shaping through energy-transition collaboration mechanisms to SAF-specific demand-side evidence.",
        "title": "Table 3: Policy as a co-evolutionary force, key concepts and SAF-specific evidence",
        "rows": [
            ("group", "Policy as business model shaper"),
            ("data", ["Policy-business model co-evolution", "Wasserbaur et al. (2022)", "Policy and business models co-evolve; neither is a fixed input to the other", "Policy shapes and is shaped by business model choices", "ReFuelEU and Dutch instruments set constraints but do not specify who builds HTL capacity or how costs are shared"]),
            ("data", ["Policy as enabler, not automatic driver", "Bryant et al. (2019)", "Government energy strategies treat policy as enabling business model change rather than guaranteeing it", "Policy alone does not create viable models", "SDL and ReFuelEU create conditions; ENLENS must still show who captures value and bears risk"]),
            ("data", ["Incumbent adaptation under policy pressure", "Leisen et al. (2019)", "German energy incumbents adapted business models unevenly under Energiewende pressure", "Incumbents adapt selectively, not uniformly", "Airports and fuel incumbents may adopt SAF procurement without circular feedstock redesign"]),
            ("group", "Energy-transition collaboration mechanisms"),
            ("data", ["Business models as system stabiliser or disruptor", "Wainstein and Bumpus (2016)", "Business models can reinforce carbon-intensive systems or enable low-carbon alternatives", "Collaboration design determines direction of change", "HTL coalition must deliberately redistribute value or reinforce existing fuel-chain logic"]),
            ("data", ["Aviation actor response to mandates", "Al-Saleh and Mahroum (2015)", "Airlines respond to SAF mandates through compliance, lobbying, or voluntary leadership depending on cost and reputation incentives", "Predicts heterogeneous airline responses at Schiphol", "KLM-style leadership is one path; minimum-compliance procurement is the default"]),
            ("group", "SAF-specific evidence"),
            ("data", ["Sector-level fare and growth constraints", "Gössling and Humpe (2023)", "Aviation's volume-growth, low-margin model is unlikely to survive heavy SAF dependence without fare restructuring", "Sets a ceiling on voluntary SAF uptake", "Even successful HTL supply may face demand limits if fare structures cannot absorb cost"]),
            ("data", ["Mandate subsidy thresholds", "Liang et al. (2026)", "SAF blending mandates require subsidies covering at least 70% of the cost differential to restore airline profitability", "Public support shapes viable blending levels", "Dutch/EU support instruments must close a similar gap for pre-commercial HTL output"]),
            ("data", ["Uneven mandate cost burden", "Karanki and Yu (2026)", "Low-cost carriers bear disproportionate SAF costs at low blending levels", "Cost pass-through is constrained by airline business model", "ReFuelEU fine caps buyer willingness to pay, limiting pass-through to passengers"]),
            ("data", ["Corporate demand-side governance", "Muller and Wittmer (2026)", "Swiss firms prefer high-blend SAF but often lack internal budget authority; WTP falls below market prices", "Internal governance, not preference, is the barrier", "Corporate voluntary demand for Schiphol-origin SAF depends on buyer-side budget structures"]),
            ("data", ["Offtake as adoption driver", "Smith et al. (2017)", "Airline offtake commitment, not technology or policy alone, drives SAF project bankability in the U.S. Pacific Northwest", "Long-term offtake precedes investment", "HTL at Schiphol needs credible airline offtake before capital deployment"]),
        ],
    },
    4: {
        "caption": "Table 4 synthesizes the circular business model design and waste valorization literature reviewed in this section, separating architectural design logic from open circular economy and valorization mechanics.",
        "title": "Table 4: Circular business model design and waste valorization, key concepts",
        "rows": [
            ("group", "Circular business model architecture"),
            ("data", ["Simultaneous circularity across components", "Geissdoerfer et al. (2018a)", "Optimal circular performance requires value proposition, creation/delivery, and capture to go circular together", "Partial circularity underperforms", "HTL must align feedstock collection, conversion, certification, and offtake simultaneously"]),
            ("data", ["Circular BMI in heavy industry", "Chirumalla et al. (2024)", "Capital-intensive manufacturing transitions require staged circular BMI with supplier and customer co-development", "Regulated heavy industry parallels aviation fuel", "Schiphol-HTL faces similar capital intensity and multi-tier supplier coordination"]),
            ("data", ["Reverse-cycle configurations", "Hansen and Revellio (2020)", "Four configurations: vertical integration, network, outsourcing, laissez-faire; asset specificity drives choice", "HTL is asset-specific, not standardised recycling", "Asset-specific HTL likely needs more vertical integration than outsourcing"]),
            ("group", "Waste valorization and open circular economy"),
            ("data", ["Open circular economy business model", "Fedele and Formisano (2023)", "Models operating across organisational boundaries with heterogeneous waste streams and shared infrastructure", "Multi-actor valorization beyond firm boundaries", "Airport waste HTL inherently requires open circular logic across operators and certifiers"]),
            ("data", ["Value-uncaptured diagnostic", "Yang et al. (2017)", "Framework identifies uncaptured value across material, energy, and capability flows in industrial systems", "Diagnostic for where value currently leaks", "Can locate where Schiphol waste value is lost before HTL conversion"]),
            ("data", ["Coproduct revenue and pathway economics", "Short et al. (2014)", "Coproduct streams materially affect biofuel pathway viability beyond primary fuel output", "Revenue diversification affects bankability", "HTL coproducts may matter for pre-commercial economics if jet yield alone is insufficient"]),
        ],
    },
    5: {
        "caption": "Table 5 synthesizes the ecosystem, coalition formation, and governance literature reviewed in this section, from multi-actor analysis through coalition dynamics to governance frameworks for circular fuel ecosystems.",
        "title": "Table 5: Ecosystem dynamics, coalition formation, and value chain governance, key concepts",
        "rows": [
            ("group", "Multi-actor and ecosystem-level analysis"),
            ("data", ["Focal-actor ecosystem design", "Hellstrom et al. (2015)", "Individual firm-level BMI is insufficient; a focal actor must manage interdependencies across the ecosystem", "Requires orchestration beyond bilateral contracts", "No obvious focal actor exists yet for Schiphol HTL; role must be negotiated"]),
            ("data", ["Shared value proposition without orchestrator", "Brink (2022)", "Offshore wind firms aligned around LCoE without a central coordinator", "Ecosystem-as-structure can coordinate without hierarchy", "A shared proposition (certified circular SAF at viable cost) could substitute for a single lead firm"]),
            ("data", ["Value network and resource pool co-evolution", "Speich and Ulli-Beer (2023)", "Unit of analysis shifts to co-evolution of value network and resource pool", "Resource pool must be designed with the network", "Schiphol waste streams, HTL tech, and certification form the resource pool"]),
            ("group", "Coalition formation"),
            ("data", ["Iterative coalition formation", "Bidmon and Knab (2018); Velter et al. (2021)", "Coalitions form iteratively under institutional pressure, power asymmetries, and boundary work; power reconfiguration is hardest", "Incumbent collaboration is necessary but fragile", "HTL operators must work with entrenched fuel incumbents rather than replace them"]),
            ("data", ["Incumbent collaboration playbook", "Wadin et al. (2017)", "Race-to-learn risk is higher in technology alliances with extractable knowledge than in business model co-development", "Alliance type shapes knowledge leakage risk", "Technology partnerships alone may not secure durable coalition logic"]),
            ("data", ["Aviation biofuel partnership precedent", "Mousavi and Bossink (2017)", "KLM's biofuel programme shows how airline-led procurement can anchor supplier investment", "Airline offtake can orchestrate upstream investment", "Provides a partial precedent, though HTL adds waste and certification complexity"]),
            ("data", ["Leadership and ecosystem maturity", "Planko et al. (2017)", "Technological innovation system functions explain why some ecosystems develop leaders and others stall", "Not every ecosystem produces a lead firm", "Schiphol may need deliberate leadership design rather than assuming emergence"]),
            ("group", "Value chain co-evolution and governance"),
            ("data", ["Material efficiency vs. emissions efficiency", "Axelson et al. (2021)", "Material efficiency strategies require greater business model change than emissions efficiency measures", "Circular fuel is closer to material efficiency", "Implies reconfiguring supply chains and actor networks, not just contracts or prices"]),
            ("data", ["Focal-actor choice shapes supply chain", "Carraresi and Broring (2021)", "Dedicated intermediary can accelerate adoption faster than incumbent core transformation", "Leadership choice is a design variable", "Relevant to whether Schiphol, an airline, or an intermediary should lead HTL"]),
            ("data", ["Governance forms in circular pioneers", "Minoja and Romano (2024)", "Hub-and-spoke, lead-role, and shared governance each distribute power differently across circular value chains", "Governance form shapes risk and value distribution", "Schiphol HTL coalition must choose among governance forms, not assume one default"]),
        ],
    },
    6: {
        "caption": "Table 6 synthesizes the financial viability, risk allocation, and design-implementation gap literature reviewed in this section, including the analytical framework derived for the ENLENS case.",
        "title": "Table 6: Financial viability, risk allocation, and the design-implementation gap, key concepts",
        "rows": [
            ("group", "Financial viability and risk allocation"),
            ("data", ["Designed risk and value distribution", "Boons and Ludeke-Freund (2013)", "Sustainable circular business models require explicit mechanisms for distributing risk and value; absence drives the design-implementation gap", "Fair distribution is a design requirement", "Without pre-designed redistribution, weaker coalition members bear disproportionate HTL costs"]),
            ("data", ["Asymmetric SAF cost burden", "Karanki and Yu (2026)", "Mandate costs fall unevenly across airline business models", "Risk allocation cannot be assumed equitable", "Cost pass-through limits at Schiphol will fall differently on network vs. low-cost carriers"]),
            ("data", ["Bankability as binding constraint", "Overholm (2017)", "Bankability must be secured before other alliance problems can be addressed; offtaker alliances are harder than financial partnerships for novel pathways", "Offtake precedes finance for HTL", "Public co-investment (Innovation Fund CfDs, SDL) can provide revenue certainty financial partners require"]),
            ("group", "Design-implementation gap"),
            ("data", ["Systematic design-implementation gap", "Geissdoerfer et al. (2018b)", "Circular BMI designs routinely fail at implementation because institutional, supply-chain, and financial conditions are missing", "Gap is structural, not managerial alone", "HTL pilot success at lab scale does not resolve implementation conditions"]),
            ("data", ["Organizational readiness for circular BMI", "Bocken et al. (2019)", "Firms need dynamic capabilities to move from promising design to operational circular models", "Capacity building is part of the pathway", "Schiphol and partners must build readiness across waste, conversion, and certification simultaneously"]),
            ("group", "Analytical framework (Section 8.5)"),
            ("data", ["Four-level analysis model", "Foss and Saebi (2017); Hellstrom et al. (2015); Konietzko et al. (2020); Boons and Ludeke-Freund (2013)", "Firm, inter-firm, ecosystem, and institutional levels interact; firm choices are constrained by higher levels", "Integrates literature reviewed in Sections 4-8", "Provides the scaffold against which Section 9 interview evidence and Section 10 discussion are read"]),
            ("data", ["ENLENS-specific literature gap", "Synthesis (Section 8.4)", "No existing study addresses policy-driven circular innovation requiring multi-actor coalition formation under certification-intensive aviation constraints", "Gap is case-specific, not generic BMI", "This review and the ENLENS project are positioned to address that gap empirically"]),
        ],
    },
}

CONCISE = {
    "bm_unit_analysis": "Business model transformation analysis requires definitional clarity. Geissdoerfer et al. (2018b) note that the field of business model innovation remains somewhat ambiguous. Across the literature, a business model is understood as the architecture of value creation, delivery, and capture (Teece, 2010). Richter (2013) applies Osterwalder and Pigneur's (2009) canvas to analyse incumbent energy-sector strategies rather than design new ones; in German utilities, incumbents developed viable models for large-scale renewables but not for distributed generation, suggesting fit with existing industry recipes may matter more than technological readiness. Foss and Saebi (2017) define business model innovation as designed, novel, nontrivial changes to key elements and/or their architecture. Adding a SAF procurement line is modular change; redesigning airport waste management as a circular feedstock supply chain is architectural change, and the literature reviewed below finds the latter more difficult.",
    "zhang_shakeel": "Zhang et al. (2024), reviewing 1,032 BMI publications (2000-2021), find growing interest in sustainable/circular BMI but limited ecosystem-level analysis, highly relevant to ENLENS where value spans waste generators, operators, certifiers, and fuel consumers. Shakeel et al. (2020) distinguish BMI (novelty in each component), SBM (sustainability embedded in components), and SBMI (both co-present), requiring coordinated change across the entire business model.",
    "liang_karanki": "Liang et al. (2026) find that SAF blending mandates require subsidies covering at least 70% of the cost differential to restore airline profitability under a 5% blend, a logic relevant beyond China. Karanki and Yu (2026) show mandate costs fall unevenly across airlines: low-cost carriers, with more price-sensitive customers and lower markups, bear disproportionate burdens at low blending levels. ReFuelEU's non-compliance fine (Article 12) caps what buyers will pay, constraining pass-through regardless of airline pricing strategy.",
    "geissdoerfer_circular": "Geissdoerfer et al. (2018a) find optimal circular performance when value proposition, creation/delivery, and capture all go circular simultaneously. Kaipainen et al. (in press) confirm this empirically: the most successful Finnish and Italian circular transitions innovated at both firm and supply-chain level, and neither alone sufficed.",
    "hellstrom_speich": "Hellstrom et al. (2015) argue individual business model design is insufficient for ecosystem-level value creation; a focal actor must identify shared opportunities and manage interdependencies. Speich and Ulli-Beer (2023) shift the unit of analysis to co-evolution of value network and resource pool. For Schiphol-HTL, the resource pool comprises quantified waste streams, HTL technology, Rotterdam-SARAS pipeline infrastructure, and certification frameworks; the value network must be designed around deploying that pool so all participants achieve viable outcomes.",
    "bidmon_velter": "Circular coalitions form iteratively, not by design. Bidmon and Knab (2018) attribute persistence to institutional pressure, power asymmetries, and cognitive barriers. Velter et al. (2021), studying a Dutch food SME, show sustainable BMI emerging through continuous boundary work on identity, power, competence, and efficiency; power reconfigurations prove most resistant. For HTL at Schiphol, operators must work with entrenched incumbents rather than replace them. Wadin et al. (2017) note race-to-learn risk is higher for technology alliances where knowledge is extractable than for business model co-development.",
    "axelson_carraresi": "Axelson et al. (2021), studying European steel decarbonisation, find material efficiency strategies (closest to circular fuel) require greater business model change than emissions efficiency measures, mainly in value creation and delivery rather than capture. For ENLENS, this implies reconfiguring supply chains and actor networks, not just contracts or prices. Carraresi and Broring (2021), on phosphate recovery, show focal-actor choice shapes supply chain design; a dedicated intermediary may accelerate adoption faster than an incumbent transforming its core business, relevant to whether Schiphol, an airline, or a circular-fuel intermediary should lead HTL.",
    "boons_karanki": "Sustainable circular business models require explicit mechanisms for distributing risk and value (Boons and Ludeke-Freund, 2013); their absence drives the design-implementation gap. Karanki and Yu (2026) show that without pre-designed redistribution, weaker participants bear disproportionate SAF costs, undermining coalition stability.",
    "certification": "Commission Implementing Regulation (EU) 2022/996 (European Commission, 2022) defines audit, mass balance, and traceability rules before SAF claims can be made. Article 10 requires on-site auditing of all economic operators delivering five or more tonnes per month (Article 13(4)); for ENLENS, Schiphol, the HTL operator, and logistics handlers each carry separate obligations. Article 19(2)(1) mandates quarterly mass balance reconciliation, adding administrative burden for pre-commercial operators and shaping coalition formation around Di Vaio et al.'s (2026) claimability requirement. Parallel EU ETS MRV converts RED III certificates into reduced surrender obligations: under the Monitoring and Reporting Regulation (European Commission, 2018), only the biomass fraction of SAF is zero-rated (Art. 38(2)), with airlines relying on supplier-certified emission factors (Art. 53(6)-(7)) rather than independent analysis.",
    "gossling": "Gössling and Humpe (2023) model four SAF cost scenarios toward net zero and find aviation's volume-growth, low-margin model is unlikely to survive heavy SAF dependence without fare restructuring; even the cheapest synthetic fuel path would require roughly 80% higher average fares by 2050, setting a ceiling on adoption regardless of supply-side innovation.",
    "muller_smith": "Muller and Wittmer (2026) find Swiss firms prefer certified high-blend SAF but often lack budget authority; corporate WTP on long-haul routes falls far below market prices, pointing to internal governance rather than preference as the barrier. Smith et al. (2017) find airline offtake commitment, not technology or policy, drives SAF adoption in the U.S. Pacific Northwest: projects become bankable only when long-term offtake agreements with credible airlines are secured.",
    "demand_side_synthesis": "Together, these studies frame SAF viability as business model- and governance-dependent rather than technical: identical technologies can succeed or fail depending on cost and risk allocation and coalition sustainability. ReFuelEU Article 5's anti-tankering obligation anchors demand geographically at Schiphol (European Parliament & Council of the European Union, 2023a, Art. 5), while Article 15's book-and-claim feasibility study (results expected Q2 2026) could widen the buyer pool or introduce competition for local producers.",
    "omgevingswet": "At Dutch level, the Omgevingswet and Besluit activiteiten leefomgeving (BAL) implement EU waste classifications. Neither LAP3 nor CMP lists HTL as a standard route for sewage sludge, so permits require case-by-case approval with higher risk and longer timelines (Ministerie van Infrastructuur en Waterstaat, n.d.). CMP classifies liquid fuel from waste as other recovery, not recycling. Where CMP chain plans remain incomplete, LAP3 sector plans still govern used oil, plastics, and animal by-products (Ministerie van Infrastructuur en Waterstaat, 2019). Schiphol's mixed waste may require compliance with both frameworks by fraction; food-derived waste needs separate NVWA authorisation under Regulation (EC) No 1069/2009.",
    "hansen": "Hansen and Revellio (2020) identify four reverse-cycle configurations in consumer electronics: vertical integration, network, outsourcing, and laissez-faire. Standardised recycling suits outsourcing; asset-specific pathways like HTL require more vertical integration than typical closed-loop activities, even though EU law classifies HTL as other recovery.",
    "brink": "Brink (2022) shows offshore wind firms aligned around a shared value proposition (LCoE) without a central coordinator, extending Adner's (2017) ecosystem-as-structure logic. For waste-to-aviation, a proposition such as certified circular SAF at competitive cost could coordinate Schiphol, airlines, and HTL operators where no single orchestrator is obvious, complementing Hellstrom et al.'s (2015) focal-firm requirement.",
    "overholm": "Overholm (2017) finds bankability is the binding constraint before other alliance problems can be addressed. For HTL, offtaker alliances are harder to build than financial partnerships, reversing the solar-PV pattern where customer alliances were straightforward. Public co-investment (EU Innovation Fund CfDs, Dutch SDL) can provide the revenue certainty financial partners require, though waste-to-SAF must still develop market demand beyond the mandate.",
    "synthesis_has": "The reviewed literature establishes that circular transformation in regulated sectors requires architectural change (Foss & Saebi, 2017), that the design-implementation gap is systematic (Geissdoerfer et al., 2018b), and that policy alone does not create business model innovation (Bryant et al., 2019). Energy-transition literature shows coordination requires collaboration mechanisms (Hellstrom et al., 2015); waste valorization literature confirms symbiosis ventures face demanding legitimacy requirements (Kanda et al., 2024); ecosystem governance literature shows circular fuel pathways need arrangements beyond bilateral contracting (Konietzko et al., 2020); and fair value distribution is a design requirement (Boons & Ludeke-Freund, 2013).",
    "synthesis_gaps": "The literature does not yet address the ENLENS case: policy-driven circular innovation requiring multi-actor coalition formation in a certification-intensive sector, with pre-commercial technology converting heterogeneous feedstocks under parallel waste-management, aviation supply-chain, and carbon-accounting constraints. Aviation economics literature treats business model as a sector-level construct (Karanki & Yu, 2026; Gössling & Humpe, 2023) rather than firm-level value architecture, and ecosystem governance frameworks (Konietzko et al., 2020; Speich & Ulli-Beer, 2023) are not grounded in aviation or circular fuel specifically.",
}


def iter_block_items(parent):
    for child in parent.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def clone_element(element):
    return copy.deepcopy(element)


def apply_concise(text):
    rules = [
        ("Approaching business model transformation analysis should start", "bm_unit_analysis"),
        ("Zhang et al. (2024), through a bibliometric review", "zhang_shakeel"),
        ("Liang et al. (2026), using system dynamics modeling", "liang_karanki"),
        ("The Geissdoerfer et al. (2018a) study created an integrative framework", "geissdoerfer_circular"),
        ("Hellstrom et al. (2015) found that individual business model design is insufficient", "hellstrom_speich"),
        ("Circular value chain coalitions", "bidmon_velter"),
        ("These requirements raise questions about which aspects of the business model", "axelson_carraresi"),
        ("Sustainable circular business models need explicit, designed mechanisms", "boons_karanki"),
        ("The certification infrastructure that connects feedstocks", "certification"),
        ("Commission Implementing Regulation (EU) 2022/996", "certification"),
        ("Good intentions may get companies onto the SAF runway", "gossling"),
        ("Muller and Wittmer (2026) find that Swiss-headquartered firms", "muller_smith"),
        ("Together, these studies outline the demand-side parameters", "demand_side_synthesis"),
        ("At the Dutch level, the Omgevingswet (Environment and Planning Act)", "omgevingswet"),
        ("Aviation, it turns out, isn't the first industry", "hansen"),
        ("Aviation, it turns out, isn\u2019t the first industry", "hansen"),
        ("Brink (2022) looked at offshore wind", "brink"),
        ("Risk redistribution alone is insufficient. Coalition formation ultimately depends", "overholm"),
        ("The reviewed literature creates substantial shoulders for the ENLENS research", "synthesis_has"),
        ("The literature does not, however, deal with the unique makeup of the ENLENS case", "synthesis_gaps"),
    ]
    for prefix, key in rules:
        if text.startswith(prefix) or prefix in text[:100]:
            return CONCISE[key]
    return text


def fix_garbled(text):
    text = text.replace("can draw from the fund und", "can draw from the fund under Article 3(3).")
    text = re.sub(r"^er Article 3\(3\)\.\s*", "", text)
    return text


def cleanup_garbled_text(text):
    text = text.replace("can draw from the fund und", "can draw from the fund under Article 3(3).")
    text = re.sub(r"^er Article 3\(3\)\.\s*", "", text)
    # Remove repeated mid-sentence fragments from the original garbled split
    text = re.sub(r"(under Article 3\(3\)\.)(?:\s*er Article 3\(3\)\.\s*)+", r"\1 ", text)
    text = re.sub(r"\.er Article 3\(3\)\.", ".", text)
    return text.strip()


def merge_policy_items(items):
    """Fix split sentence across policy paragraphs and garbled fund/Article 3(3) text."""
    texts = []
    tables = []
    for item in items:
        if item[0] == "tbl":
            tables.append(item)
        else:
            texts.append(item[1])

    merged = []
    i = 0
    while i < len(texts):
        t = texts[i]
        while i + 1 < len(texts) and texts[i + 1].strip().startswith("er Article 3(3)"):
            nxt = re.sub(r"^er Article 3\(3\)\.\s*", "", texts[i + 1].strip())
            if t.rstrip().endswith("und") and "fund" in t:
                t = t.rstrip()[:-3] + "under Article 3(3). " + nxt
            elif "under Article 3(3)" in t:
                t = t.rstrip() + " " + nxt
            else:
                t = t.rstrip() + " " + nxt
            i += 1
        merged.append(cleanup_garbled_text(t))
        i += 1

    out = [("p", t, None) for t in merged if t]
    out.extend(tables)
    return out


def split_bryant(text):
    marker = "At the Dutch level, the Klimaatwet brings"
    if marker in text:
        idx = text.index(marker)
        return text[:idx].strip(), text[idx:].strip()
    return text, None


def extract_wasserbaur_theory(text):
    marker = "The SDL program is financed through"
    if marker in text:
        theory = text[: text.index(marker)].strip()
        policy = text[text.index(marker) :].strip()
        # Remove incomplete sentence fragment left at end of theory part
        theory = re.sub(r", and can draw from the fund und$", ".", theory)
        theory = re.sub(r", and can draw from the fund under Article 3\(3\)\.$", ".", theory)
        return theory, policy
    return text, None


ZW_CHAR = "\u200b\u200c\u200d\ufeff"
ZW = rf"[{ZW_CHAR}]?"


def heading_prefix_from(doc_or_para):
    zw_re = re.compile(rf"^([{ZW_CHAR}])\d+\.")
    if hasattr(doc_or_para, "paragraphs"):
        for p in doc_or_para.paragraphs:
            m = zw_re.match(p.text)
            if m:
                return m.group(1)
        return ""
    m = re.match(rf"^([{ZW_CHAR}])", doc_or_para.text)
    return m.group(1) if m else ""

def is_top_heading(text):
    return bool(re.match(rf"^{ZW}(\d+)\.\s+[A-Za-z]", text))


def section_num(text):
    m = re.match(rf"^{ZW}(\d+)\.", text)
    return int(m.group(1)) if m else None


def renumber_section_items(items, old, new):
    out = []
    for item in items:
        if item[0] == "p":
            t = re.sub(rf"^({ZW}){old}\.", rf"\g<1>{new}.", item[1])
            if t == item[1]:
                t = re.sub(rf"^{old}\.", f"{new}.", item[1])
            out.append(("p", apply_concise(t), item[2]))
        else:
            out.append(item)
    return out


def replace_para_text(p, text):
    """Replace paragraph text while keeping first-run formatting."""
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in list(p.runs[1:]):
        r._element.getparent().remove(r._element)
    pPr = p._element.find(qn("w:pPr"))
    if pPr is not None:
        for shd in pPr.findall(qn("w:shd")):
            pPr.remove(shd)


def strip_highlight_from_para(p):
    pPr = p._element.find(qn("w:pPr"))
    if pPr is not None:
        for shd in pPr.findall(qn("w:shd")):
            pPr.remove(shd)
    for r in p.runs:
        rPr = r._element.find(qn("w:rPr"))
        if rPr is not None:
            for shd in rPr.findall(qn("w:shd")):
                rPr.remove(shd)


def para_is_centered(p):
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        return False
    jc = pPr.find(qn("w:jc"))
    return jc is not None and jc.get(qn("w:val")) == "center"


def apply_para_style_from_template(target_p, template_p):
    """Copy paragraph and first-run character properties from a template."""
    old_pPr = target_p._element.find(qn("w:pPr"))
    if old_pPr is not None:
        target_p._element.remove(old_pPr)
    src_pPr = template_p._element.find(qn("w:pPr"))
    if src_pPr is not None:
        target_p._element.insert(0, clone_element(src_pPr))

    if not target_p.runs:
        target_p.add_run("")
    if template_p.runs:
        old_rPr = target_p.runs[0]._element.find(qn("w:rPr"))
        if old_rPr is not None:
            target_p.runs[0]._element.remove(old_rPr)
        src_rPr = template_p.runs[0]._element.find(qn("w:rPr"))
        if src_rPr is not None:
            target_p.runs[0]._element.insert(0, clone_element(src_rPr))


def set_tc_text(tc, text):
    """Replace text in a table cell while keeping run/paragraph XML structure."""
    text = "" if text is None else str(text)
    ps = tc.findall(qn("w:p"))
    if not ps:
        return
    for extra in ps[1:]:
        tc.remove(extra)
    p = ps[0]
    t_nodes = p.findall(".//" + qn("w:t"))
    if t_nodes:
        t_nodes[0].text = text
        for node in t_nodes[1:]:
            node.text = ""
        return
    rs = p.findall(qn("w:r"))
    if rs:
        t = rs[0].find(qn("w:t"))
        if t is None:
            t = rs[0].makeelement(qn("w:t"), {})
            rs[0].append(t)
        t.text = text


def normalize_table_row(values):
    values = list(values)
    if len(values) == 4:
        return [values[0], values[1], values[2], "", values[3]]
    while len(values) < 5:
        values.append("")
    return values[:5]


def build_table_element(template_table, row_specs):
    """Build a table by cloning row XML from the policy Table 1 template."""
    tbl_el = clone_element(template_table._element)
    old_rows = tbl_el.findall(qn("w:tr"))
    for tr in old_rows:
        tbl_el.remove(tr)

    header_tr = clone_element(template_table.rows[0]._tr)
    group_tr = template_table.rows[1]._tr
    data_tr = template_table.rows[2]._tr

    set_tc_texts_in_tr(header_tr, LIT_TABLE_HEADERS)
    tbl_el.append(header_tr)

    for spec in row_specs:
        if spec[0] == "group":
            tr = clone_element(group_tr)
            label = spec[1]
            set_tc_texts_in_tr(tr, [label] * 5, group=True)
        else:
            tr = clone_element(data_tr)
            set_tc_texts_in_tr(tr, normalize_table_row(spec[1]))
        tbl_el.append(tr)
    return tbl_el


def set_tc_texts_in_tr(tr, values, group=False):
    tcs = tr.findall(qn("w:tc"))
    if group:
        label = values[0] if values else ""
        for tc in tcs:
            set_tc_text(tc, label)
        return
    for i, tc in enumerate(tcs):
        set_tc_text(tc, values[i] if i < len(values) else "")


def build_interview_table_element(template_table, rows):
    """Four-column interview table using Table 1 styling."""
    tbl_el = clone_element(template_table._element)
    for tr in tbl_el.findall(qn("w:tr")):
        tbl_el.remove(tr)

    header_tr = clone_element(template_table.rows[0]._tr)
    data_tr = template_table.rows[2]._tr
    set_tc_texts_in_tr(header_tr, [rows[0][0], rows[0][1], rows[0][2], rows[0][3], ""])
    tbl_el.append(header_tr)
    for row in rows[1:]:
        tr = clone_element(data_tr)
        set_tc_texts_in_tr(tr, [row[0], row[1], row[2], row[3], ""])
        tbl_el.append(tr)
    return tbl_el


def restore_ooxml_parts(orig_path, out_path, parts):
    """Put selected OOXML parts back from the original after python-docx save."""
    with zipfile.ZipFile(orig_path) as z_orig:
        orig_data = {part: z_orig.read(part) for part in parts}
    buf = io.BytesIO()
    with zipfile.ZipFile(out_path, "r") as z_in, zipfile.ZipFile(buf, "w") as z_out:
        for item in z_in.infolist():
            data = orig_data[item.filename] if item.filename in orig_data else z_in.read(item.filename)
            z_out.writestr(item, data)
    with open(out_path, "wb") as f:
        f.write(buf.getvalue())


def find_templates(doc):
    title_t = doc.paragraphs[0] if doc.paragraphs else None
    empty_t = doc.paragraphs[1] if len(doc.paragraphs) > 1 else title_t
    body_t = heading_t = ref_t = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if not heading_t and t.startswith("1. Introduction"):
            heading_t = p
        if not body_t and t.startswith("Unlike most industrial sectors"):
            body_t = p
        if (
            not body_t
            and len(t) > 100
            and p.runs
            and not para_is_centered(p)
            and not t.startswith("Business Model Transformation")
        ):
            body_t = p
        if not ref_t and t.startswith("Adner") and "2017" in t:
            ref_t = p
    if not body_t and len(doc.paragraphs) > 37:
        body_t = doc.paragraphs[37]
    if not heading_t:
        heading_t = body_t
    if not ref_t:
        for p in doc.paragraphs:
            if p.text.strip().startswith("References"):
                idx = doc.paragraphs.index(p)
                if idx + 1 < len(doc.paragraphs):
                    ref_t = doc.paragraphs[idx + 1]
                    break
    if not ref_t:
        ref_t = body_t
    return title_t, empty_t, body_t, heading_t, ref_t


def main():
    shutil.copy2(SRC, OUT)
    src = Document(SRC)
    blocks = list(iter_block_items(src))

    title_text = None
    sections = {}
    refs = []
    current = None
    in_refs = False

    for b in blocks:
        if isinstance(b, Paragraph):
            t = b.text.strip()
            if in_refs:
                refs.append(("p", t, b))
                continue
            if t.startswith("References"):
                in_refs = True
                refs.append(("p", t, b))
                current = None
                continue
            sn = section_num(t) if is_top_heading(t) else None
            if sn is not None:
                current = sn
                sections.setdefault(current, []).append(("p", t, b))
            elif current is None:
                if t and not t.startswith("1) Methodology") and not t.startswith("-") and title_text is None:
                    title_text = t
            else:
                sections.setdefault(current, []).append(("p", t, b))
        else:
            if in_refs:
                refs.append(("tbl", b))
            elif current is not None:
                sections.setdefault(current, []).append(("tbl", b))

    # Split section 3 into policy vs theory
    policy_items = []
    theory_items = []
    pending_policy = []

    for item in sections.get(3, []):
        if item[0] == "tbl":
            policy_items.append(item)
            continue
        t = fix_garbled(item[1])
        if t.startswith("3. Regulatory Pressure"):
            theory_items.append(("p", "5. Literature: Policy as a Co-Evolutionary Force", item[2]))
        elif t.startswith("3.1 Policy as Business Model Shaper"):
            theory_items.append(("p", "5.1 Policy as Business Model Shaper: A Discussion of Co-evolutionary Dynamics", item[2]))
        elif t.startswith("3.2 Co-evolutionary"):
            continue
        elif t.startswith("3.3"):
            theory_items.append(("p", "5.3 Lessons from the Energy Transition: Collaboration Mechanisms and the Chicken-and-Egg Problem", item[2]))
        elif t.startswith("3.4"):
            theory_items.append(("p", "5.4 SAF-Specific Evidence on Aviation Business Model Pressure", item[2]))
        elif "Table 1 synthesizes" in t or t.startswith("Table 1: Key regulatory"):
            policy_items.append(("p", t, item[2]))
        elif t.startswith("The certification infrastructure") or "Commission Implementing Regulation (EU) 2022/996" in t[:60]:
            policy_items.append(("p", apply_concise(t), item[2]))
        elif "RED III Article 27" in t or ("Time is of the essence" in t and "HTL project" in t):
            policy_items.append(("p", t, item[2]))
        elif t.strip().startswith("er Article 3(3)"):
            policy_items.append(("p", t, item[2]))
        elif "Before turning to the Dutch and Schiphol-specific case" in t and "co-evolution" in t:
            theory_part, policy_tail = extract_wasserbaur_theory(t)
            theory_items.append(("p", theory_part, item[2]))
            if policy_tail:
                pending_policy.append(policy_tail)
        elif t.startswith("In contrast, Bryant et al."):
            theory_part, policy_part = split_bryant(t)
            theory_items.append(("p", apply_concise(theory_part), item[2]))
            if policy_part:
                pending_policy.append(policy_part)
        elif t.startswith("At the Dutch level, the Omgevingswet"):
            policy_items.append(("p", apply_concise(t), item[2]))
        else:
            theory_items.append(("p", apply_concise(t), item[2]))

    for p in pending_policy:
        policy_items.insert(0, ("p", fix_garbled(p), None))

    policy_items = merge_policy_items(policy_items)

    # Move Omgevingswet from section 4 to policy
    sec4 = []
    for item in sections.get(4, []):
        if item[0] == "p" and item[1].startswith("At the Dutch level, the Omgevingswet"):
            policy_items.append(("p", apply_concise(item[1]), item[2]))
        else:
            sec4.append(item)
    sections[4] = sec4

    sec2 = renumber_section_items(sections.get(2, []), 2, 4)
    sec6 = renumber_section_items(sections.get(4, []), 4, 6)
    sec7 = renumber_section_items(sections.get(5, []), 5, 7)
    sec8 = renumber_section_items(sections.get(6, []), 6, 8)
    sec8_synth = renumber_section_items(sections.get(7, []), 7, 9)
    # Fold synthesis into section 8
    synth_folded = []
    for item in sec8_synth:
        if item[0] == "p":
            t = item[1]
            t = t.replace("9. Synthesis: Gaps, Contributions, and Analytical Framework", "")
            t = t.replace("9.1 What the Literature Has and Has Not Addressed", "8.4 What the Literature Has and Has Not Addressed")
            t = t.replace("9.2 The Analytical Framework", "8.5 The Analytical Framework")
            if t.strip():
                synth_folded.append(("p", apply_concise(t), item[2]))
        else:
            synth_folded.append(item)

    with open("_new_sections.txt") as f:
        raw = f.read()
    interviews = raw.split("===INTERVIEWS===\n")[1].split("\n===DISCUSSION===\n")[0].strip()
    discussion = raw.split("===DISCUSSION===\n")[1].split("\n===CONCLUSION===\n")[0].strip()
    conclusion = raw.split("===CONCLUSION===\n")[1].strip()

    # Rebuild output from copy of original, cloning paragraph/table formatting
    dst = Document(OUT)
    body = dst.element.body
    title_template, empty_template, body_template, heading_template, ref_template = find_templates(src)
    table_template = src.tables[0] if src.tables else None
    heading_prefix = heading_prefix_from(src)
    table_caption_template = body_template
    table_title_template = body_template
    for p in src.paragraphs:
        t = p.text.strip()
        if t.startswith("Table 1 synthesizes"):
            table_caption_template = p
        if t.startswith("Table 1: Key regulatory"):
            table_title_template = p

    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    def is_heading_text(text):
        t = text.strip()
        return bool(re.match(rf"^{ZW}\d+(\.\d+)*\.\s+\S", t)) or re.match(r"^Table \d+:", t) or t.startswith("Table 1:")

    def add_table_title(text):
        new_el = clone_element(table_title_template._element)
        body.append(new_el)
        p = Paragraph(new_el, dst)
        replace_para_text(p, str(text).strip())
        strip_highlight_from_para(p)
        return p

    def add_table_caption(text):
        new_el = clone_element(table_caption_template._element)
        body.append(new_el)
        p = Paragraph(new_el, dst)
        replace_para_text(p, str(text).strip())
        strip_highlight_from_para(p)
        return p

    def add_overview_table(table_num):
        spec = OVERVIEW_TABLES[table_num]
        add_table_caption(spec["caption"])
        add_table_title(spec["title"])
        body.append(build_table_element(table_template, spec["rows"]))
        add_para("")

    def add_interview_table(rows):
        add_table_title("Interview participants")
        body.append(build_interview_table_element(table_template, rows))
        add_para("")

    def format_heading(text):
        t = str(text).strip()
        if is_heading_text(t) and heading_prefix and not t.startswith(heading_prefix):
            return heading_prefix + t
        return t

    def add_para(text, template=None):
        template = template or body_template
        new_el = clone_element(template._element)
        body.append(new_el)
        p = Paragraph(new_el, dst)
        cleaned = str(text).strip() if text else ""
        if template is heading_template:
            cleaned = format_heading(cleaned)
        replace_para_text(p, cleaned)
        strip_highlight_from_para(p)
        return p

    def add_table_from_source(source_table):
        body.append(clone_element(source_table._element))

    def add_para_from_item(item):
        if item[0] == "tbl":
            add_table_from_source(item[1])
            return
        new_text = item[1]
        orig_para = item[2] if len(item) > 2 else None
        if orig_para is not None:
            new_el = clone_element(orig_para._element)
            body.append(new_el)
            p = Paragraph(new_el, dst)
            text_to_set = format_heading(new_text) if is_heading_text(new_text) else new_text
            if text_to_set.strip() != orig_para.text.strip():
                replace_para_text(p, text_to_set)
                strip_highlight_from_para(p)
            return
        tmpl = heading_template if is_heading_text(new_text) else body_template
        add_para(new_text, tmpl)

    if title_text:
        add_para(title_text, title_template)
    add_para("", empty_template)

    add_para("1. Introduction", heading_template)
    for block in INTRO.split("\n\n"):
        add_para(block.strip())

    add_para(f"{heading_prefix}2. Methodology", heading_template)
    add_para(METHODOLOGY)
    add_para(f"{heading_prefix}2.1 Literature search strategy and databases", heading_template)
    add_para(METH_21)
    add_para(f"{heading_prefix}2.2 Deduplication and screening", heading_template)
    add_para(METH_22)
    add_para(f"{heading_prefix}2.3 Reference checking", heading_template)
    add_para(METH_23)
    add_para(f"{heading_prefix}2.4 Regulatory and legal instrument mapping", heading_template)
    add_para(METH_24)
    add_para(f"{heading_prefix}2.5 Interview approach", heading_template)
    add_para(METH_25)
    interview_rows = [
        ["Participant", "Position in value chain", "Organisation type", "Interview focus"],
        ["Strategy department representative", "Infrastructure and logistics, non-fuel-transacting", "Port authority", "Storage, blending, terminal investment, coordination role of infrastructure actors"],
        ["M&A and holdings lead, board member of a SAF joint venture", "Fuel buyer and mandated compliance party", "Airline", "Offtake contracting, project finance, internal capability building, cost pass-through"],
        ["Former partnership manager, emission-reduction solutions", "Certificate and voluntary-market intermediary", "SAF producer (former employee)", "Book-and-claim mechanics, corporate voluntary demand, certification standards"],
        ["Head of partnerships and technology development", "Upstream technology supplier", "Catalyst manufacturer", "Feedstock and pathway diversification, R&D portfolio allocation, cross-sector spillovers"],
        ["Head of partnerships and technology, low carbon fuels", "Producer, trader, and project developer", "Integrated energy major", "Risk allocation across the value chain, offtake and financing mismatch, regulatory design"],
    ]
    add_interview_table(interview_rows)

    add_para(f"{heading_prefix}3. Policy Context", heading_template)
    add_para(POLICY_OPENING)
    for item in policy_items:
        add_para_from_item(item)

    for item in sec2:
        add_para_from_item(item)
    add_overview_table(2)

    for item in theory_items:
        add_para_from_item(item)
    add_overview_table(3)

    for item in sec6:
        add_para_from_item(item)
    add_overview_table(4)

    for item in sec7:
        add_para_from_item(item)
    add_overview_table(5)

    for item in sec8:
        add_para_from_item(item)
    for item in synth_folded:
        add_para_from_item(item)
    add_overview_table(6)

    for block in interviews.split("\n\n"):
        if block.strip():
            tmpl = heading_template if is_heading_text(block) else body_template
            add_para(format_heading(block.strip()) if tmpl is heading_template else block.strip(), tmpl)
    for block in discussion.split("\n\n"):
        if block.strip():
            tmpl = heading_template if is_heading_text(block) else body_template
            add_para(format_heading(block.strip()) if tmpl is heading_template else block.strip(), tmpl)
    for block in conclusion.split("\n\n"):
        if block.strip():
            tmpl = heading_template if is_heading_text(block) else body_template
            add_para(format_heading(block.strip()) if tmpl is heading_template else block.strip(), tmpl)

    add_para(f"{heading_prefix}References", heading_template)
    for item in refs[1:]:
        if item[0] == "p":
            new_el = clone_element(ref_template._element)
            body.append(new_el)
            p = Paragraph(new_el, dst)
            replace_para_text(p, item[1])

    dst.save(OUT)
    restore_ooxml_parts(SRC, OUT, OOXML_PARTS_TO_RESTORE)

    # Final pass: clean garbled fragments, cross-refs, section numbering, em dashes
    final = Document(OUT)
    refs_idx = next(
        (i for i, p in enumerate(final.paragraphs) if p.text.strip() == "References"),
        len(final.paragraphs),
    )
    replacements = [
        ("5.3 Lessons from the Energy Transition", "5.2 Lessons from the Energy Transition"),
        ("5.4 SAF-Specific Evidence on Aviation Business Model Pressure", "5.3 SAF-Specific Evidence on Aviation Business Model Pressure"),
        ("Sections 5 and 6 of the literature review", "Sections 6 and 7 of the literature review"),
        ("the analytical framework in Section 8.2 has not yet been tested", "the analytical framework in Section 8.5 has not yet been tested"),
        ("Section 5.4, including Gössling", "Section 5.3, including Gössling"),
        (
            "described in Section 2.5 and summarised in Table 1.",
            "described in Section 2.5 and summarised in the interview participants table in Section 2.",
        ),
    ]
    for i, p in enumerate(final.paragraphs):
        text = cleanup_garbled_text(p.text)
        for old, new in replacements:
            text = text.replace(old, new)
        if i < refs_idx:
            text = text.replace("—", ", ").replace("–", ", ")
        if text != p.text:
            replace_para_text(p, text)
    final.save(OUT)
    restore_ooxml_parts(SRC, OUT, OOXML_PARTS_TO_RESTORE)

    print(f"Saved {OUT}")


if __name__ == "__main__":
    main()
